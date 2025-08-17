"""负责将聊天记录导出为 Word 或 PDF。"""

from __future__ import annotations

import tempfile
from datetime import datetime
from typing import List, Dict
import os
import pdfkit
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# -------------------- 内部工具 --------------------

def _get_wkhtmltopdf_path() -> str:
    """获取 wkhtmltopdf 可执行文件路径。抄自旧实现。"""
    local_path = os.path.join(os.path.dirname(__file__), "..", "bin", "wkhtmltopdf.exe")
    if os.path.exists(local_path):
        return local_path
    if os.system("where wkhtmltopdf > nul 2>&1") == 0:
        return "wkhtmltopdf"
    program_files_path = r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"
    if os.path.exists(program_files_path):
        return program_files_path
    raise FileNotFoundError("未找到wkhtmltopdf，请确保已安装或放置在正确位置")

# PDF 选项常量
PDF_OPTIONS = {
    "page-size": "A4",
    "margin-top": "20mm",
    "margin-right": "20mm",
    "margin-bottom": "20mm",
    "margin-left": "20mm",
    "encoding": "UTF-8",
    "custom-header": [("Accept-Encoding", "gzip")],
    "no-outline": None,
    "quiet": "",
}

# -------------------- Word 生成 --------------------

def _generate_word(messages: List[Dict[str, str]], title: str, temp_file_path: str) -> None:
    doc = Document()
    # 页面 & 样式
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Cm(2.54)

    title_style = doc.styles.add_style("CustomTitle", WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = "微软雅黑"
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_style.paragraph_format.space_after = Pt(20)

    doc.add_paragraph(title, style="CustomTitle")

    time_paragraph = doc.add_paragraph()
    time_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    time_run = time_paragraph.add_run(f"创建时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    time_run.font.name = "微软雅黑"
    time_run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    time_run.font.size = Pt(10)
    time_run.font.color.rgb = RGBColor(128, 128, 128)

    # 分隔线
    separator = doc.add_paragraph()
    separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    separator_run = separator.add_run("=" * 40)
    separator_run.font.size = Pt(12)
    separator_run.font.color.rgb = RGBColor(200, 200, 200)

    # 消息样式
    msg_style = doc.styles.add_style("MessageStyle", WD_STYLE_TYPE.PARAGRAPH)
    msg_style.font.name = "微软雅黑"
    msg_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    msg_style.font.size = Pt(11)
    msg_style.paragraph_format.space_before = Pt(12)
    msg_style.paragraph_format.space_after = Pt(12)
    msg_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for msg in messages:
        role_para = doc.add_paragraph(style="MessageStyle")
        role_run = role_para.add_run(f"{msg['role']}：")
        role_run.bold = True
        content_para = doc.add_paragraph(style="MessageStyle")
        content_para.add_run(msg["content"])
        doc.add_paragraph(style="MessageStyle")  # 空行间隔

    doc.save(temp_file_path)

# -------------------- PDF 生成 --------------------

def _generate_pdf(messages: List[Dict[str, str]], title: str, temp_file_path: str) -> None:
    html_parts = [
        "<html><head><meta charset='utf-8'><style>",
        "@font-face {font-family: 'Microsoft YaHei'; src: local('Microsoft YaHei');}",
        "body {font-family:'Microsoft YaHei', Arial, sans-serif; margin:40px; line-height:1.6;}",
        ".title {text-align:center;font-size:24px;margin-bottom:10px;font-weight:bold;}",
        ".time {text-align:center;font-size:12px;color:#666;margin-bottom:20px;}",
        ".divider {border-top:1px solid #ccc;margin:20px 0;}",
        ".role {font-weight:bold;margin-top:15px;color:#000;}",
        ".content {margin:5px 0 15px 20px;color:#333;}",
        "</style></head><body>",
        f"<div class='title'>{title}</div>",
        f"<div class='time'>创建时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>",
        "<div class='divider'></div>",
    ]
    for msg in messages:
        html_parts.append(f"<div class='role'>{msg['role']}：</div>")
        html_parts.append(f"<div class='content'>{msg['content']}</div>")
    html_parts.append("</body></html>")
    html_content = "".join(html_parts)

    wkhtmltopdf_path = _get_wkhtmltopdf_path()
    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
    pdfkit.from_string(html_content, temp_file_path, options=PDF_OPTIONS, configuration=config)

# -------------------- Public API --------------------

def generate_export(messages: List[Dict[str, str]], title: str, fmt: str) -> str:
    """生成导出文件并返回文件路径。"""
    if fmt not in {"word", "pdf"}:
        raise ValueError("不支持的导出格式")

    suffix = ".docx" if fmt == "word" else ".pdf"
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    temp_file.close()  # 我们将手动写入

    if fmt == "word":
        _generate_word(messages, title, temp_file.name)
    else:
        _generate_pdf(messages, title, temp_file.name)

    return temp_file.name 