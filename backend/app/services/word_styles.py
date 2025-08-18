from __future__ import annotations

"""封装 Word 样式逻辑，供 export_service 复用。"""

from datetime import datetime
from typing import List, Dict

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def create_document(title: str) -> Document:
    """创建带默认页面设置与标题的 Document。"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Cm(2.54)

    # 标题样式
    title_style = doc.styles.add_style("CustomTitle", WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = "微软雅黑"
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_style.paragraph_format.space_after = Pt(20)

    doc.add_paragraph(title, style="CustomTitle")

    time_para = doc.add_paragraph()
    time_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    time_run = time_para.add_run(f"创建时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    time_run.font.name = "微软雅黑"
    time_run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    time_run.font.size = Pt(10)
    time_run.font.color.rgb = RGBColor(128, 128, 128)

    divider = doc.add_paragraph()
    divider.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    divider_run = divider.add_run("=" * 40)
    divider_run.font.size = Pt(12)
    divider_run.font.color.rgb = RGBColor(200, 200, 200)

    # 消息样式
    msg_style = doc.styles.add_style("Msg", WD_STYLE_TYPE.PARAGRAPH)
    msg_style.font.name = "微软雅黑"
    msg_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    msg_style.font.size = Pt(11)
    msg_style.paragraph_format.space_before = Pt(12)
    msg_style.paragraph_format.space_after = Pt(12)
    msg_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    return doc


def append_messages(doc: Document, messages: List[Dict[str, str]]) -> None:
    """把聊天消息追加到文档。"""
    for msg in messages:
        role_para = doc.add_paragraph(style="Msg")
        role_para.add_run(f"{msg['role']}：").bold = True
        content_para = doc.add_paragraph(style="Msg")
        content_para.add_run(msg["content"])
        doc.add_paragraph(style="Msg")  # 空行 